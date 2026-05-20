"""Claude-generated equity research report (.docx).

Calls Anthropic's Claude API with the user's existing prompt (prompt_consensus.txt)
plus structured financial data. Writes the result as a formatted Word document.

Also returns a structured dict of "one_pager_blocks" that the PPTX generator consumes.
"""
from __future__ import annotations
import json
import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..extract.income_statement import IncomeStatement
from ..model.dcf import DCFResult
from ..data.yfinance_client import QuoteSnapshot

load_dotenv(override=True)


@dataclass
class ReportPayload:
    """Structured output from Claude, parsed from JSON envelope.

    New 6-section Deep Research Report format:
      1. executive_summary — recommendation, thesis, fair value, key strengths/risks
      2. business_model     — how they make money (string)
      3. moat_competition   — moat + competitors + tech advantage (dict)
      4. catalysts_12mo     — list of upcoming catalysts (list of dicts)
      5. risk_assessment    — skeptic's view: accounting / concentration / competitive (dict)
      6. asymmetry_check    — valuation floor vs growth ceiling, verdict (dict)
    """
    executive_summary: Dict[str, object] = field(default_factory=dict)
    business_model: str = ""
    moat_competition: Dict[str, object] = field(default_factory=dict)
    catalysts_12mo: List[Dict[str, str]] = field(default_factory=list)
    risk_assessment: Dict[str, str] = field(default_factory=dict)
    asymmetry_check: Dict[str, str] = field(default_factory=dict)
    # One-pager components (consumed by pptx generator) — unchanged
    one_pager: Dict[str, object] = field(default_factory=dict)
    # Raw markdown in case parsing failed
    raw_markdown: str = ""


def _financial_context(stmt: IncomeStatement, dcf: DCFResult, quote: QuoteSnapshot) -> str:
    """Render compact financial context for Claude."""
    lines: List[str] = []
    lines.append(f"# {stmt.title} ({stmt.ticker})")
    lines.append(f"Current price: ${quote.current_price}")
    if quote.market_cap:
        lines.append(f"Market cap: ${quote.market_cap/1e9:.2f}B")
    if quote.enterprise_value:
        lines.append(f"Enterprise value: ${quote.enterprise_value/1e9:.2f}B")
    lines.append(f"Sector: {quote.sector} / Industry: {quote.industry}")
    lines.append("")
    lines.append("## Quarterly Income Statement (USD thousands)")
    period_labels = [p.label for p in stmt.periods]
    lines.append("Period | " + " | ".join(period_labels))
    for line in stmt.lines:
        row = [line.label]
        is_eps = line.section == "eps"
        for p in period_labels:
            v = line.values.get(p)
            if v is None:
                row.append("-")
            elif is_eps:
                row.append(f"{v:.2f}")
            else:
                row.append(f"{v/1000:,.0f}")
        lines.append(" | ".join(row))
    lines.append("")

    lines.append("## DCF Valuation Output (Bank-Style)")
    lines.append(f"Base year: {dcf.base_fy}, base revenue: ${dcf.base_revenue/1e9:.2f}B")
    lines.append(f"Methodology: {dcf.assumptions.explicit_years}Y explicit + {dcf.assumptions.fade_years}Y fade + dual terminal value (Gordon Growth + Exit Multiple)")
    w = dcf.wacc_base
    lines.append(f"WACC (CAPM): {w.wacc:.2%}  (Ke={w.cost_of_equity:.2%}, Kd={w.aftertax_cost_of_debt:.2%} after-tax, E/V={w.equity_weight:.0%}, β={w.inputs.beta:.2f})")
    lines.append(f"Terminal growth: {dcf.assumptions.terminal_growth:.1%},  Peer median EV/EBITDA: {dcf.peer_median_ev_ebitda:.1f}x" if dcf.peer_median_ev_ebitda else "")
    lines.append("")
    lines.append("Scenario outputs (blended Gordon Growth + Exit Multiple):")
    for name in ("Bear", "Base", "Bull"):
        s = dcf.scenarios[name]
        upside = (s.price_per_share / dcf.current_price - 1) if dcf.current_price else None
        up_str = f" ({upside:+.0%} vs current)" if upside is not None else ""
        lines.append(f"  {name}: WACC {s.wacc:.2%}, EV ${s.enterprise_value/1e9:.2f}B, Equity ${s.equity_value/1e9:.2f}B, Price ${s.price_per_share:.2f}{up_str}")
    lines.append("")

    # Analyst consensus context
    if dcf.analyst_count or dcf.analyst_target_mean:
        lines.append("## Analyst Consensus (sell-side)")
        if dcf.analyst_count:
            lines.append(f"Analyst coverage: {dcf.analyst_count} analysts, rating = {dcf.analyst_rec_key or 'n/a'}")
        if dcf.analyst_target_mean:
            lines.append(f"Price targets: low ${dcf.analyst_target_low:.2f} / mean ${dcf.analyst_target_mean:.2f} / high ${dcf.analyst_target_high:.2f}")
        if dcf.analyst_growth_y1 is not None:
            lines.append(f"Consensus revenue growth: Y1 {dcf.analyst_growth_y1:+.1%}, Y2 {dcf.analyst_growth_y2:+.1%}")
            lines.append(f"(DCF projections seeded from these consensus estimates, then fade to terminal growth.)")
        lines.append("")

    lines.append("## TTM Market Fundamentals")
    for attr in ("revenue_ttm", "ebitda_ttm", "free_cash_flow", "pe_ratio", "forward_pe",
                 "ev_ebitda", "ev_sales", "net_leverage"):
        v = getattr(quote, attr, None)
        if v is not None:
            if isinstance(v, float) and abs(v) > 1e6:
                lines.append(f"{attr}: ${v/1e9:.2f}B")
            else:
                lines.append(f"{attr}: {v}")

    return "\n".join(lines)


def _build_system_prompt(prompt_template: str) -> str:
    """Combine user's research prompt with structured-output instructions."""
    return f"""{prompt_template}

---

OUTPUT FORMAT REQUIREMENTS:

Return your response as a single JSON object wrapped in <report> tags. Use this exact schema:

<report>
{{
  "executive_summary": {{
    "recommendation": "BUY" | "HOLD" | "AVOID" | "WAIT FOR [CONDITIONS]",
    "thesis": "2-3 sentence investment thesis",
    "current_price": "$XX.XX",
    "fair_value": "$XX.XX",
    "upside_downside_pct": "+XX%" | "-XX%",
    "conviction": "High" | "Medium" | "Low",
    "key_strengths": ["3 short bullets, factual"],
    "key_risks": ["3 short bullets, factual"],
    "action_today": "1 sentence: buy at current / wait for pullback to $X / avoid entirely"
  }},

  "business_model": "2-3 paragraphs explaining how the company makes money in plain English. Who pays, for what, how often, at what price. Include revenue stream mix if you can infer it from the data.",

  "moat_competition": {{
    "moat_description": "1-2 paragraphs on the unique advantage (technology, patents, network effects, scale, regulatory exclusivity). If they don't have a real moat, say so explicitly.",
    "top_competitors": [
      {{"name": "Competitor name", "ticker": "TICKER or N/A", "comparison": "1-2 sentences on how they stack up vs the target"}},
      {{"name": "Competitor 2", "ticker": "...", "comparison": "..."}},
      {{"name": "Competitor 3", "ticker": "...", "comparison": "..."}}
    ],
    "tech_advantage": "Specific patents, IP, proprietary tech, or other defensible advantage. State 'No meaningful tech moat' if there isn't one."
  }},

  "catalysts_12mo": [
    {{"event": "Event description", "timing": "Q1 2026 / H2 2026 / etc", "impact": "High" | "Medium" | "Low", "rationale": "Why it matters to the stock"}},
    "...3 to 5 entries"
  ],

  "risk_assessment": {{
    "accounting_irregularities": "1 paragraph as a skeptic: aggressive revenue recognition, SBC masking losses, off-balance-sheet items, restatements, audit issues, related-party transactions. Cite specific 10-K/10-Q line items if any look concerning. If clean, say so.",
    "concentration_risk": "1 paragraph: customer concentration, single-product dependency, geographic concentration, single-platform/channel risk, regulatory single-point-of-failure.",
    "competitive_threats": "1 paragraph: emerging entrants, technology shifts, pricing pressure, substitution risk. Reference well-known short theses for the company where they exist (drawing on common bear arguments — you don't have live web access)."
  }},

  "asymmetry_check": {{
    "valuation_floor": "1 paragraph on what protects the downside: hard assets, cash position, recurring revenue, scarcity value, contracted backlog. Quantify where possible.",
    "growth_ceiling": "1 paragraph on the realistic upside scenario: TAM expansion, margin recovery, new product categories, geographic growth.",
    "verdict": "Asymmetric to upside" | "Roughly symmetric" | "Asymmetric to downside",
    "reasoning": "1-2 paragraphs synthesizing valuation, growth, and the risks from section 4. State the verdict clearly."
  }},

  "one_pager": {{
    "tagline": "One-line company descriptor, e.g. 'AI Global Learning Platform'",
    "intro_paragraph": "2-3 sentences summarizing the company for the one-pager (150 words max)",
    "operations_footprint": ["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"],
    "key_strengths": ["bullet 1", "bullet 2", "bullet 3", "bullet 4", "bullet 5"],
    "key_risks": ["bullet 1", "bullet 2", "bullet 3", "bullet 4"],
    "primary_kpi_label": "Paid Subscribers" | "MAU" | "Transacting Users" | etc,
    "primary_kpi_series": {{}}
  }}
}}
</report>

Write in the tone of an institutional equity research memo — analytical, thesis-driven, data-heavy. Cite specific numbers from the provided financial context. Do not use emoji.
"""


def call_claude(stmt: IncomeStatement, dcf: DCFResult, quote: QuoteSnapshot,
                prompt_template: str, model: str = "claude-sonnet-4-6",
                max_tokens: int = 8000, temperature: float = 0.4) -> ReportPayload:
    """Call Claude and parse the structured response."""
    try:
        import anthropic
    except ImportError as e:
        raise RuntimeError("anthropic SDK not installed") from e

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise RuntimeError("ANTHROPIC_API_KEY not set in environment or .env file")

    client = anthropic.Anthropic(api_key=api_key)

    context = _financial_context(stmt, dcf, quote)
    system = _build_system_prompt(prompt_template)
    user_msg = (f"Analyze the following company using the framework above.\n\n"
                f"{context}\n\n"
                f"Generate the complete report in the specified JSON schema.")

    msg = client.messages.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        system=system,
        messages=[{"role": "user", "content": user_msg}],
    )

    text = "".join(block.text for block in msg.content if hasattr(block, "text"))
    return _parse_claude_response(text)


def _parse_claude_response(text: str) -> ReportPayload:
    """Extract the JSON envelope from Claude's response."""
    payload = ReportPayload(raw_markdown=text)
    match = re.search(r"<report>\s*(\{.*?\})\s*</report>", text, re.DOTALL)
    if not match:
        # Try to find a raw JSON block
        match = re.search(r"\{.*\"executive_summary\".*\}", text, re.DOTALL)
    if not match:
        return payload
    json_str = match.group(1) if match.re.pattern.startswith("<report>") else match.group(0)
    try:
        data = json.loads(json_str)
    except json.JSONDecodeError:
        return payload
    payload.executive_summary = data.get("executive_summary", {})
    payload.business_model = data.get("business_model", "")
    payload.moat_competition = data.get("moat_competition", {})
    payload.catalysts_12mo = data.get("catalysts_12mo", []) or []
    payload.risk_assessment = data.get("risk_assessment", {})
    payload.asymmetry_check = data.get("asymmetry_check", {})
    payload.one_pager = data.get("one_pager", {})
    return payload


# ---------- Word document writer ----------

def _set_paragraph_font(para, size=11, bold=False, color=None, italic=False):
    for run in para.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def _add_hr(doc):
    """Add a thin horizontal rule (paragraph with bottom border)."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "A6A6A6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc, text: str, size: int = 14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3864")
    _add_hr(doc)


def _add_body_paragraph(doc, text: str):
    if not text:
        return
    # Split by single newline to preserve paragraph structure
    for para_text in text.split("\n\n"):
        if not para_text.strip():
            continue
        p = doc.add_paragraph()
        run = p.add_run(para_text.strip())
        run.font.size = Pt(11)


def _add_bullet_list(doc, items: List[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def _add_kv_block(doc, pairs: List[tuple[str, str]]):
    """Add a key-value block for executive summary metrics."""
    for k, v in pairs:
        p = doc.add_paragraph()
        bold = p.add_run(f"{k}: ")
        bold.font.bold = True
        bold.font.size = Pt(11)
        val = p.add_run(str(v))
        val.font.size = Pt(11)


def _build_relval_table(quote: QuoteSnapshot, peer_quotes: List[QuoteSnapshot]) -> List[Dict]:
    """Compute the 3-row relative-valuation table: target + top 2 peers.

    Columns: P/S TTM, P/S Forward, EV/EBITDA, Gross Margin, YoY Revenue Growth,
    Value/Growth Score (P/S TTM ÷ revenue growth %).
    Lower V/G score = more growth per dollar of valuation.
    """
    def row(q: QuoteSnapshot) -> Dict[str, object]:
        ps_ttm = (q.market_cap / q.revenue_ttm) if (q.market_cap and q.revenue_ttm) else None
        ps_fwd = (q.market_cap / q.analyst_rev_current_year) if (q.market_cap and q.analyst_rev_current_year) else None
        gross_margin = (q.gross_profit_ttm / q.revenue_ttm) if (q.gross_profit_ttm and q.revenue_ttm) else None
        # YoY revenue growth: prefer consensus current-year growth, else fallback to None
        yoy_growth = q.analyst_rev_growth_current_year
        # Value/Growth score = P/S TTM ÷ (growth % × 100). Lower = better.
        vg_score = None
        if ps_ttm and yoy_growth and yoy_growth > 0:
            vg_score = ps_ttm / (yoy_growth * 100)
        return {
            "ticker": q.ticker,
            "name": q.company_name or q.ticker,
            "ps_ttm": ps_ttm,
            "ps_fwd": ps_fwd,
            "ev_ebitda": q.ev_ebitda,
            "gross_margin": gross_margin,
            "yoy_growth": yoy_growth,
            "vg_score": vg_score,
        }

    rows = [row(quote)]
    # Preserve the curated order from peers.yaml / sector defaults — the first
    # peers in the list are typically the most narratively relevant. Skip any
    # without valid market_cap data.
    relevant_peers = [p for p in (peer_quotes or []) if p and p.market_cap]
    for p in relevant_peers[:2]:
        rows.append(row(p))
    return rows


def _add_relval_table(doc: "Document", rows: List[Dict]):
    """Insert the relative-valuation table into the Word doc."""
    if not rows:
        return
    headers = ["Ticker", "P/S TTM", "P/S Fwd", "EV/EBITDA", "Gross Margin", "YoY Rev Growth", "Value/Growth Score"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("FFFFFF")

    # Data rows
    for i, r in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = r["ticker"]
        cells[1].text = f"{r['ps_ttm']:.2f}x" if r['ps_ttm'] else "—"
        cells[2].text = f"{r['ps_fwd']:.2f}x" if r['ps_fwd'] else "—"
        cells[3].text = f"{r['ev_ebitda']:.1f}x" if r['ev_ebitda'] else "—"
        cells[4].text = f"{r['gross_margin']*100:.0f}%" if r['gross_margin'] else "—"
        cells[5].text = f"{r['yoy_growth']*100:+.1f}%" if r['yoy_growth'] is not None else "—"
        cells[6].text = f"{r['vg_score']:.3f}" if r['vg_score'] else "—"
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                # First column (ticker) bold
                if cell == cells[0]:
                    for run in p.runs:
                        run.font.bold = True

    # Caption
    cap = doc.add_paragraph()
    cap_run = cap.add_run("Lower Value/Growth Score = more growth per dollar of valuation. Computed from yfinance fundamentals + analyst consensus.")
    cap_run.font.italic = True
    cap_run.font.size = Pt(8)
    cap_run.font.color.rgb = RGBColor.from_string("595959")


def write_docx(payload: ReportPayload, stmt: IncomeStatement, quote: QuoteSnapshot,
               output_path: Path,
               peer_quotes: Optional[List[QuoteSnapshot]] = None) -> Path:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    run = title.add_run(f"{stmt.ticker}  —  {stmt.title}")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3864")

    subtitle = doc.add_paragraph()
    run = subtitle.add_run("Deep Research Report  ·  Investment Committee Memo")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor.from_string("595959")

    _add_hr(doc)

    # ===== Executive Summary =====
    _add_section_heading(doc, "Executive Summary")
    es = payload.executive_summary
    if es:
        _add_kv_block(doc, [
            ("Recommendation", es.get("recommendation", "-")),
            ("Thesis", es.get("thesis", "-")),
            ("Current Price", es.get("current_price", f"${quote.current_price or '-'}")),
            ("Fair Value", es.get("fair_value", "-")),
            ("Upside / Downside", es.get("upside_downside_pct", "-")),
            ("Conviction", es.get("conviction", "-")),
            ("Action Today", es.get("action_today", "-")),
        ])
        if es.get("key_strengths"):
            doc.add_paragraph().add_run("Key Strengths").font.bold = True
            _add_bullet_list(doc, es.get("key_strengths", []))
        if es.get("key_risks"):
            doc.add_paragraph().add_run("Key Risks").font.bold = True
            _add_bullet_list(doc, es.get("key_risks", []))

    # ===== 1. Business Model =====
    _add_section_heading(doc, "1.  Business Model")
    _add_body_paragraph(doc, payload.business_model)

    # ===== 2. Moat & Competition =====
    _add_section_heading(doc, "2.  Moat & Competition")
    mc = payload.moat_competition or {}
    _add_body_paragraph(doc, mc.get("moat_description", ""))

    # Competitor list
    comps = mc.get("top_competitors", []) or []
    if comps:
        doc.add_paragraph().add_run("Top Competitors").font.bold = True
        for c in comps:
            p = doc.add_paragraph(style="List Bullet")
            name_run = p.add_run(f"{c.get('name', 'Unknown')}")
            name_run.font.bold = True
            name_run.font.size = Pt(11)
            ticker_str = c.get('ticker', '')
            if ticker_str and ticker_str != "N/A":
                tr = p.add_run(f" ({ticker_str})")
                tr.font.size = Pt(11)
            txt_run = p.add_run(f" — {c.get('comparison', '')}")
            txt_run.font.size = Pt(11)

    # Tech advantage
    tech = mc.get("tech_advantage", "")
    if tech:
        doc.add_paragraph().add_run("Technological / Defensible Advantage").font.bold = True
        _add_body_paragraph(doc, tech)

    # Relative valuation table (computed deterministically from yfinance)
    if peer_quotes:
        rv_rows = _build_relval_table(quote, peer_quotes)
        if rv_rows:
            doc.add_paragraph().add_run("Relative Valuation Snapshot").font.bold = True
            _add_relval_table(doc, rv_rows)

    # ===== 3. Catalysts =====
    _add_section_heading(doc, "3.  Catalysts (Next 12 Months)")
    cats = payload.catalysts_12mo or []
    if cats:
        for c in cats:
            p = doc.add_paragraph(style="List Bullet")
            event = p.add_run(c.get("event", "Event"))
            event.font.bold = True
            event.font.size = Pt(11)
            timing = c.get("timing", "")
            impact = c.get("impact", "")
            meta = []
            if timing:
                meta.append(timing)
            if impact:
                meta.append(f"Impact: {impact}")
            if meta:
                meta_run = p.add_run(f"  ({' · '.join(meta)})")
                meta_run.font.size = Pt(10)
                meta_run.font.italic = True
                meta_run.font.color.rgb = RGBColor.from_string("595959")
            rationale = c.get("rationale", "")
            if rationale:
                r2 = p.add_run(f"\n    {rationale}")
                r2.font.size = Pt(10)

    # ===== 4. Risk Assessment =====
    _add_section_heading(doc, "4.  Risk Assessment  (skeptic's view)")
    ra = payload.risk_assessment or {}
    if ra.get("accounting_irregularities"):
        doc.add_paragraph().add_run("Accounting / Reporting").font.bold = True
        _add_body_paragraph(doc, ra.get("accounting_irregularities", ""))
    if ra.get("concentration_risk"):
        doc.add_paragraph().add_run("Customer / Geographic / Segment Concentration").font.bold = True
        _add_body_paragraph(doc, ra.get("concentration_risk", ""))
    if ra.get("competitive_threats"):
        doc.add_paragraph().add_run("Competitive Threats").font.bold = True
        _add_body_paragraph(doc, ra.get("competitive_threats", ""))

    # ===== 5. Asymmetry Check =====
    _add_section_heading(doc, "5.  Asymmetry Check")
    ac = payload.asymmetry_check or {}
    if ac.get("valuation_floor"):
        doc.add_paragraph().add_run("Valuation Floor (Downside Protection)").font.bold = True
        _add_body_paragraph(doc, ac.get("valuation_floor", ""))
    if ac.get("growth_ceiling"):
        doc.add_paragraph().add_run("Growth Ceiling (Upside Scenario)").font.bold = True
        _add_body_paragraph(doc, ac.get("growth_ceiling", ""))
    verdict = ac.get("verdict", "")
    if verdict:
        p = doc.add_paragraph()
        run = p.add_run(f"Verdict:  {verdict}")
        run.font.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor.from_string("1F3864")
    if ac.get("reasoning"):
        _add_body_paragraph(doc, ac.get("reasoning", ""))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def write_fallback_docx(stmt: IncomeStatement, quote: QuoteSnapshot, note: str,
                         output_path: Path) -> Path:
    """If Claude API is unavailable, write a skeleton doc with the data context and a note."""
    doc = Document()
    title = doc.add_paragraph()
    run = title.add_run(f"{stmt.ticker}  —  {stmt.title}")
    run.font.size = Pt(20); run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F3864")

    p = doc.add_paragraph()
    r = p.add_run(note)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("C00000")

    doc.add_paragraph()
    doc.add_paragraph("Financial context (to paste into ChatGPT if you prefer manual generation):")
    from src.model.dcf import DCFResult  # type: ignore
    # Just the context string, no DCF passed
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def write_prompt_file(stmt: IncomeStatement, dcf: DCFResult, quote: QuoteSnapshot,
                      prompt_template: str, output_path: Path) -> Path:
    """Write the exact prompt + context to a .txt for manual iteration."""
    context = _financial_context(stmt, dcf, quote)
    system = _build_system_prompt(prompt_template)
    content = f"=== SYSTEM PROMPT ===\n\n{system}\n\n=== USER MESSAGE ===\n\nAnalyze the following company using the framework above.\n\n{context}\n"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")
    return output_path
